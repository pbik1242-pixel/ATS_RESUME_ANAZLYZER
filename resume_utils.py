import base64
import io
import os
import re
import uuid
import logging
from pathlib import Path
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor

# Optional OCR / PDF-to-image imports (only used if available)
try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    import pytesseract
    _HAS_OCR = True
except Exception:
    _HAS_OCR = False
_HAS_PDF_PREVIEW = convert_from_path is not None


_EMBED = None
_EMBED_UTIL = None
_EMBED_ATTEMPTED = False


def _configure_pdf_logging() -> None:
    # pdfplumber/pdfminer can emit noisy non-fatal font metadata warnings for
    # some resumes. Keep request logs and real exceptions visible, but suppress
    # these parser warnings in normal app output.
    for logger_name in ("pdfminer", "pdfminer.pdffont", "pdfminer.pdfinterp", "pdfplumber"):
        logging.getLogger(logger_name).setLevel(logging.ERROR)


_configure_pdf_logging()


def _env_truthy(name: str) -> bool:
    return os.getenv(name, "").strip().lower() in {"1", "true", "yes", "on"}


def _restore_env(key: str, value: str | None) -> None:
    if value is None:
        os.environ.pop(key, None)
    else:
        os.environ[key] = value


def _get_embedder():
    global _EMBED, _EMBED_UTIL, _EMBED_ATTEMPTED

    if _EMBED_ATTEMPTED:
        return _EMBED, _EMBED_UTIL

    _EMBED_ATTEMPTED = True

    if not _env_truthy("ENABLE_SEMANTIC_SCORE"):
        return None, None

    try:
        from sentence_transformers import SentenceTransformer, util
    except Exception as e:
        print("Semantic model disabled:", e)
        return None, None

    model_name = os.getenv("SEMANTIC_MODEL", "all-MiniLM-L6-v2").strip() or "all-MiniLM-L6-v2"
    allow_download = _env_truthy("ALLOW_MODEL_DOWNLOAD")

    orig_hf_offline = os.environ.get("HF_HUB_OFFLINE")
    orig_tf_offline = os.environ.get("TRANSFORMERS_OFFLINE")
    try:
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        _EMBED = SentenceTransformer(model_name)
        _EMBED_UTIL = util
        return _EMBED, _EMBED_UTIL
    except Exception as e_offline:
        if not allow_download:
            print("Semantic model disabled:", e_offline)
            return None, None

        try:
            _restore_env("HF_HUB_OFFLINE", orig_hf_offline)
            _restore_env("TRANSFORMERS_OFFLINE", orig_tf_offline)
            _EMBED = SentenceTransformer(model_name)
            _EMBED_UTIL = util
            return _EMBED, _EMBED_UTIL
        except Exception as e_online:
            print("Semantic model disabled:", e_online)
            return None, None
    finally:
        _restore_env("HF_HUB_OFFLINE", orig_hf_offline)
        _restore_env("TRANSFORMERS_OFFLINE", orig_tf_offline)




def extract_text(path: Path) -> str:
    try:
        ext = path.suffix.lower()

        if ext == ".txt":
            return path.read_text(errors="ignore")

        if ext == ".pdf":
            text = ""
            with pdfplumber.open(path) as pdf:
                for p in pdf.pages:
                    t = p.extract_text()
                    if t:
                        text += t + "\n"
            if text.strip():
                return text
            # fallback to OCR if no text extracted and OCR libs available
            if _HAS_OCR:
                ocr_chunks = []
                try:
                    for img in convert_from_path(str(path)):
                        ocr_chunks.append(pytesseract.image_to_string(img))
                    return "\n".join(ocr_chunks)
                except Exception:
                    pass
            return text

        if ext == ".docx":
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)

        return ""
    except Exception:
        return ""


def _clean_cell(value: str | None) -> str:
    if value is None:
        return ""
    text = re.sub(r"\s+", " ", str(value)).strip()
    parts = text.split()
    if (
        len(parts) >= 2
        and all(part.isalnum() for part in parts)
        and (
            (len(parts) == 2 and (len(parts[0]) > 4 or len(parts[1]) <= 3))
            or (len(parts) == 3 and any(len(part) <= 2 for part in parts))
        )
    ):
        text = "".join(parts)
    return text


def _table_rows_from_pdf(path: Path) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables() or []
            for table in page_tables:
                rows: list[list[str]] = []
                for row in table or []:
                    cleaned = [_clean_cell(cell) for cell in (row or [])]
                    if any(cleaned):
                        rows.append(cleaned)
                if rows:
                    tables.append(rows)
    return tables


def _table_rows_from_docx(path: Path) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    doc = Document(path)
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cleaned = [_clean_cell(cell.text) for cell in row.cells]
            if any(cleaned):
                rows.append(cleaned)
        if rows:
            tables.append(rows)
    return tables


def extract_tables(path: Path) -> list[list[list[str]]]:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return _table_rows_from_pdf(path)
        if ext == ".docx":
            return _table_rows_from_docx(path)
        return []
    except Exception:
        return []


def build_preview_image_data(path: Path) -> str | None:
    try:
        if path.suffix.lower() != ".pdf" or not _HAS_PDF_PREVIEW:
            return None
        pages = convert_from_path(str(path), first_page=1, last_page=1, dpi=180)
        if not pages:
            return None
        buffer = io.BytesIO()
        pages[0].save(buffer, format="PNG")
        encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
        return f"data:image/png;base64,{encoded}"
    except Exception:
        return None


def tables_to_text(tables: list[list[list[str]]]) -> str:
    lines: list[str] = []
    for table in tables:
        for row in table:
            cleaned = []
            for cell in row:
                value = _clean_cell(cell)
                if value:
                    cleaned.append(value)
            if cleaned:
                lines.append(" | ".join(cleaned))
    return "\n".join(lines)


def clean(text):
    return re.sub(r"\s+", " ", text).strip()


def normalize_upload_name(filename: str | None) -> tuple[str, str]:
    raw = str(filename or "").replace("\\", "/").split("/")[-1]
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", raw).strip("._")
    safe = safe or "resume.txt"
    return safe, Path(safe).suffix.lower()


def fallback_resume_points(text: str, limit: int = 12) -> list[str]:
    src = str(text or "")
    if not src.strip():
        return []
    out: list[str] = []
    seen: set[str] = set()
    for raw in src.splitlines():
        line = re.sub(r"\s+", " ", raw).strip().lstrip("-* ").strip()
        if 35 <= len(line) <= 260 and line.lower() not in seen:
            seen.add(line.lower())
            out.append(line)
            if len(out) >= limit:
                return out
    for part in re.split(r"(?<=[.!?])\s+", clean(src)):
        part = part.strip()
        if 35 <= len(part) <= 260 and part.lower() not in seen:
            seen.add(part.lower())
            out.append(part)
            if len(out) >= limit:
                return out
    return out[:limit]


ACTION_VERBS = (
    "analyzed", "built", "created", "delivered", "designed", "developed", "drove", "executed",
    "implemented", "improved", "led", "managed", "optimized", "resolved", "streamlined", "supported",
    "worked", "owned", "launched", "collaborated", "mentored", "coordinated",
)


def _normalize_point(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "").strip())
    return cleaned.rstrip(";:,. ") + ("" if not cleaned else ".")


def combine_points(points: list[str]) -> list[str]:
    seen: set[str] = set()
    combined: list[str] = []
    for raw in points or []:
        cleaned = _normalize_point(raw)
        key = cleaned.lower()
        if len(cleaned) < 8 or key in seen:
            continue
        seen.add(key)
        combined.append(cleaned)
    return combined


def build_suggestions(points: list[str], missing: list[str]) -> list[str]:
    suggestions: list[str] = []
    missing_clean = [m.strip() for m in (missing or []) if m and m.strip() and m.strip() != "-"]
    if missing_clean:
        suggestions.append(
            "Add these missing keywords where they are true: "
            + ", ".join(missing_clean[:8])
            + ("..." if len(missing_clean) > 8 else "")
        )
    if points and any(not re.search(r"\d", p) for p in points):
        suggestions.append("Quantify impact with numbers (time saved, volume processed, % improvements).")
    if points and any(len(p) > 170 for p in points):
        suggestions.append("Shorten long bullets to 1 line and move extra detail into tools or scope lines.")
    if points and any(not re.match(r"(?i)^(" + "|".join(ACTION_VERBS) + r")\b", p) for p in points):
        suggestions.append("Start each bullet with a strong action verb to show ownership and outcomes.")
    if not points:
        suggestions.append("Add 4-6 focused bullets that describe impact, tools, and measurable results.")
    return suggestions


SKILL_LABELS = {
    "c": "C",
    "cpp": "C++",
    "cplus": "C++",
    "csharp": "C#",
    "java": "Java",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "python": "Python",
    "ruby": "Ruby",
    "php": "PHP",
    "go": "Go",
    "golang": "Go",
    "rust": "Rust",
    "kotlin": "Kotlin",
    "swift": "Swift",
    "scala": "Scala",
    "r": "R",
    "sql": "SQL",
    "mysql": "MySQL",
    "postgresql": "PostgreSQL",
    "mongodb": "MongoDB",
    "redis": "Redis",
    "oracle": "Oracle",
    "sqlite": "SQLite",
    "dotnet": ".NET",
    "net": ".NET",
    "node": "Node.js",
    "react": "React",
    "angular": "Angular",
    "vue": "Vue",
    "django": "Django",
    "flask": "Flask",
    "fastapi": "FastAPI",
    "spring": "Spring",
    "aws": "AWS",
    "azure": "Azure",
    "gcp": "GCP",
    "docker": "Docker",
    "kubernetes": "Kubernetes",
    "terraform": "Terraform",
    "git": "Git",
    "linux": "Linux",
    "powershell": "PowerShell",
    "ssis": "SSIS",
    "ssrs": "SSRS",
    "ssms": "SSMS",
}


SKILL_ALLOWLIST = set(SKILL_LABELS.keys())
SKILL_PHRASES = {
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "data analysis": "Data Analysis",
    "data analytics": "Data Analytics",
    "data engineering": "Data Engineering",
    "data science": "Data Science",
    "power bi": "Power BI",
    "tableau": "Tableau",
    "excel": "Excel",
    "etl": "ETL",
    "api": "API",
    "rest api": "REST API",
    "microservices": "Microservices",
    "pandas": "Pandas",
    "numpy": "NumPy",
    "scikit learn": "Scikit-learn",
    "tensorflow": "TensorFlow",
    "pytorch": "PyTorch",
    "html": "HTML",
    "css": "CSS",
    "bootstrap": "Bootstrap",
}


def _normalize_skill_token(token: str) -> str:
    key = re.sub(r"[^a-z0-9+.#]", "", token.lower())
    if key in SKILL_LABELS:
        return SKILL_LABELS[key]
    return token.strip()


def filter_skill_terms(terms: list[str]) -> list[str]:
    filtered: list[str] = []
    seen: set[str] = set()
    for raw in terms or []:
        token = str(raw or "").strip()
        if not token:
            continue
        key = re.sub(r"[^a-z0-9+.#]", "", token.lower())
        # keep experience years like "3 years", "5 yrs", "8+ years"
        if re.search(r"\b\d+\s*\+?\s*(years?|yrs?)\b", token, re.I):
            label = re.sub(r"\s+", " ", token)
        elif key in SKILL_ALLOWLIST:
            label = _normalize_skill_token(key)
        else:
            continue
        out_key = label.lower()
        if out_key in seen:
            continue
        seen.add(out_key)
        filtered.append(label)
    return filtered


def extract_profile_skills(text: str, limit: int = 20) -> list[str]:
    source = str(text or "")
    if not source.strip():
        return []
    lowered = source.lower()
    found: list[str] = []
    seen: set[str] = set()

    for phrase, label in SKILL_PHRASES.items():
        pattern = r"(?<![a-z0-9])" + re.escape(phrase) + r"(?![a-z0-9])"
        if re.search(pattern, lowered):
            key = label.lower()
            if key not in seen:
                seen.add(key)
                found.append(label)
                if len(found) >= limit:
                    return found

    for raw_key, label in SKILL_LABELS.items():
        pattern = r"(?<![a-z0-9])" + re.escape(raw_key) + r"(?![a-z0-9])"
        if re.search(pattern, lowered):
            key = label.lower()
            if key not in seen:
                seen.add(key)
                found.append(label)
                if len(found) >= limit:
                    return found

    return found


def build_table_preview(
    tables: list[list[list[str]]],
    max_tables: int = 2,
    max_rows: int = 6,
    max_cols: int = 6,
    max_cell_len: int = 80,
) -> list[list[list[str]]]:
    preview: list[list[list[str]]] = []
    for table in tables[:max_tables]:
        rows: list[list[str]] = []
        for row in table[:max_rows]:
            trimmed: list[str] = []
            for cell in row[:max_cols]:
                value = (cell or "").strip()
                if max_cell_len and len(value) > max_cell_len:
                    value = value[: max_cell_len - 3] + "..."
                trimmed.append(value)
            if any(trimmed):
                rows.append(trimmed)
        if rows:
            preview.append(rows)
    return preview





def extract_bullets(text):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    bullets = []

    for l in lines:
        if l.startswith(("-", "*")):
            bullets.append(l[1:].strip())
        elif len(l) < 180 and re.search(
            r"(developed|built|created|managed|designed|led|improved|worked)",
            l,
            re.I
        ):
            bullets.append(l)

    if not bullets:
        parts = re.split(r"[.!?]", text)
        for p in parts:
            if 30 < len(p) < 250:
                bullets.append(p.strip())

    return bullets


def extract_keywords(text, top_n=20):
    words = re.findall(r"[a-zA-Z]+", text.lower())
    stopwords = {"the", "and", "for", "with", "this", "that", "were", "your"}
    freq = {}

    for w in words:
        if len(w) > 2 and w not in stopwords:
            freq[w] = freq.get(w, 0) + 1

    sorted_words = sorted(freq.items(), key=lambda x: -x[1])
    return [w for w, _ in sorted_words[:top_n]]


def keyword_score(resume, job_keywords):
    found, missing = [], []
    resume_lower = resume.lower()

    for kw in job_keywords:
        if kw.lower() in resume_lower:
            found.append(kw)
        else:
            missing.append(kw)

    score = len(found) / max(1, len(job_keywords))
    return score, found, missing


def formatting_score(text):
    score = 0
    tl = text.lower()

    if "experience" in tl:
        score += 0.3
    if "education" in tl:
        score += 0.3
    if "skills" in tl:
        score += 0.3
    if len(text) < 5000:
        score += 0.1

    return min(score, 1)


def semantic_score(resume, job) -> float | None:
    embed, util = _get_embedder()
    if embed is None:
        return None
    try:
        r = embed.encode(resume, convert_to_tensor=True)
        j = embed.encode(job, convert_to_tensor=True)
        score = float(util.cos_sim(r, j).item())
        if score < 0.0:
            return 0.0
        if score > 1.0:
            return 1.0
        return score
    except Exception:
        return None



def _apply_layout(doc: Document, layout: str) -> None:
    layout = (layout or "classic").lower()
    if layout == "modern":
        base_font = ("Arial", Pt(11), RGBColor(15, 23, 42))
        accent = RGBColor(29, 78, 216)
    elif layout == "minimal":
        base_font = ("Times New Roman", Pt(11), RGBColor(33, 37, 41))
        accent = RGBColor(90, 90, 90)
    else:  # classic
        base_font = ("Calibri", Pt(11), RGBColor(31, 41, 55))
        accent = RGBColor(37, 99, 235)

    # Normal text
    normal = doc.styles["Normal"].font
    normal.name = base_font[0]
    normal.size = base_font[1]
    normal.color.rgb = base_font[2]

    for heading in ("Heading 1", "Heading 2", "Heading 3"):
        h_style = doc.styles[heading].font
        h_style.name = base_font[0]
        h_style.size = Pt(16 if heading == "Heading 1" else 13)
        h_style.color.rgb = accent


def export_docx(name, skills, bullets, export_dir, original_text: str | None = None, layout: str = "classic"):
    try:
        export_dir = Path(export_dir)
        export_dir.mkdir(parents=True, exist_ok=True)
        skills_list = [str(s).strip() for s in (skills or []) if str(s).strip()]
        bullets_list = [str(b).strip() for b in (bullets or []) if str(b).strip()]

        doc = Document()
        _apply_layout(doc, layout)
        doc.add_heading(name or "Candidate", level=0)

        if original_text:
            doc.add_heading("Original Resume Excerpt", level=1)
            # keep it concise
            snippet = original_text[:3000]
            doc.add_paragraph(snippet if snippet else "N/A")

        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills_list) if skills_list else "N/A")

        doc.add_heading("Experience", level=1)
        if bullets_list:
            for b in bullets_list:
                try:
                    doc.add_paragraph(b, style="List Bullet")
                except Exception:
                    doc.add_paragraph(f"- {b}")
        else:
            doc.add_paragraph("No bullet points available.")

        doc.add_heading("AI-Improved Points", level=1)
        if bullets_list:
            for idx, b in enumerate(bullets_list, start=1):
                doc.add_paragraph(f"{idx}. {b}")
        else:
            doc.add_paragraph("No improved points available.")

        out = export_dir / f"{uuid.uuid4().hex}.docx"
        doc.save(str(out))
        return out
    except Exception as e:
        print("DOCX export failed:", e)
        return None


def extract_contact_details(text: str, fallback_name: str = "Candidate") -> dict[str, str]:
    """Pull a best-effort name, email, and phone number from resume text.

    - Email: first RFC-ish email pattern
    - Contact: first phone-like pattern with at least 10 digits
    - Name: first non-empty line that isn't the email/phone and doesn't look like metadata
    """

    blob = text or ""
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", blob)
    email = email_match.group(0).strip() if email_match else ""

    phone_match = re.search(r"(?:\+?\d[\d\s().-]{8,}\d)", blob)
    contact = ""
    if phone_match:
        raw_phone = phone_match.group(0)
        digits = re.sub(r"[^0-9+]", "", raw_phone)
        # normalize leading 00 to + where sensible
        if digits.startswith("00") and len(digits) > 2:
            digits = "+" + digits[2:]
        contact = digits or raw_phone.strip()

    lines = [clean(l) for l in str(blob).splitlines() if clean(l)]
    candidate_name = fallback_name
    for line in lines[:8]:
        lower_line = line.lower()
        if email and email.lower() in lower_line:
            continue
        if contact and any(ch.isdigit() for ch in line):
            continue
        # avoid obvious headers
        if len(line) < 3 or len(line) > 60:
            continue
        candidate_name = line
        break

    return {
        "candidate_name": candidate_name,
        "contact": contact,
        "email": email,
    }
