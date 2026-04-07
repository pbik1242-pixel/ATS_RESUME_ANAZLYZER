from pathlib import Path

from docx import Document

from resume_utils import extract_tables, extract_text, tables_to_text


def _make_docx_with_table(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Sample Resume")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Advanced"
    doc.save(str(path))


def test_extract_tables_from_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "resume.docx"
    _make_docx_with_table(docx_path)
    tables = extract_tables(docx_path)
    assert len(tables) == 1
    assert tables[0][0][0] == "Skill"
    assert tables[0][1][1] == "Advanced"


def test_tables_to_text_renders_rows() -> None:
    tables = [[["Skill", "Level"], ["Python", "Advanced"]]]
    text = tables_to_text(tables)
    assert "Skill | Level" in text
    assert "Python | Advanced" in text


def test_extract_text_docx_includes_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Experience")
    doc.add_paragraph("Built APIs")
    doc.save(str(docx_path))
    text = extract_text(docx_path)
    assert "Experience" in text
    assert "Built APIs" in text
